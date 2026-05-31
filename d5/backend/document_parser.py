from pathlib import Path
from typing import List
import re
import chardet
from langchain.schema import Document
from langchain_community.document_loaders import TextLoader
from docx import Document as DocxDocument
import pdfplumber


class DocumentParser:
    @staticmethod
    def parse(file_path: str) -> List[Document]:
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == '.pdf':
            return DocumentParser._parse_pdf(file_path)
        elif suffix == '.docx':
            return DocumentParser._parse_docx(file_path)
        elif suffix == '.txt':
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    @staticmethod
    def _parse_pdf(file_path: str) -> List[Document]:
        documents = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    text = DocumentParser._clean_text(text)
                    if text.strip():
                        documents.append(Document(
                            page_content=text,
                            metadata={"source": file_path, "page": page_num}
                        ))
        return documents

    @staticmethod
    def _clean_text(text: str) -> str:
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\x00-\x7F\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    @staticmethod
    def _parse_docx(file_path: str) -> List[Document]:
        doc = DocxDocument(file_path)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(' | '.join(row_text))
        
        text = '\n'.join(full_text)
        return [Document(page_content=text, metadata={"source": file_path, "page": 0})]

    @staticmethod
    def _parse_txt(file_path: str) -> List[Document]:
        with open(file_path, 'rb') as f:
            raw_data = f.read()
            result = chardet.detect(raw_data)
            encoding = result['encoding'] or 'utf-8'
        
        try:
            loader = TextLoader(file_path, encoding=encoding)
            return loader.load()
        except UnicodeDecodeError:
            loader = TextLoader(file_path, encoding='gbk')
            return loader.load()
